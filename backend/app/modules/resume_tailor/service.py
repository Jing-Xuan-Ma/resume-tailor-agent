"""
Resume Tailor Service — orchestrates the LangGraph agent.
"""

from io import BytesIO
import re
from textwrap import wrap
from typing import Optional
from uuid import UUID

from app import db
from app.core.models import ParsedJobDescription, Resume, TailoredResume
from app.core.events import ResumeTailoredEvent, event_bus
from app.memory.experience_embedder import ExperienceEmbedder
from app.memory.long_term import LongTermMemoryStore
from app.modules.resume_tailor.agent import tailor_agent
from app.modules.resume_tailor.draft_store import draft_store
from app.modules.resume_tailor.nodes.file_parser import parse_resume_file
from app.modules.resume_tailor.nodes.parse_jd import JDParsingNode
from app.modules.resume_tailor.nodes.tailor_resume import TailorResumeNode
from app.modules.resume_tailor.nodes.text_export import TextExportNode


class ResumeTailorService:
    """
    High-level service for resume tailoring operations.
    """

    def __init__(self):
        self.jd_parser = JDParsingNode()
        self.embedder = ExperienceEmbedder()
        self.memory_store = LongTermMemoryStore()
        self.text_exporter = TextExportNode()
        self.tailor_node = TailorResumeNode()

    def _rebuild_resume_data(self, user_id: str) -> dict:
        """Rebuild resume_data dict from Chroma experience documents."""
        docs = self.memory_store.get_all_experiences(user_id)
        if not docs:
            return {}

        raw_doc = next(
            (doc for doc in docs if (doc.get("metadata") or {}).get("chunk_type") == "raw_resume"),
            None,
        )
        if raw_doc:
            return self._parse_standard_resume_text(raw_doc.get("text", ""))

        experiences: dict[str, dict] = {}
        for doc in docs:
            meta = doc.get("metadata", {})
            exp_id = meta.get("experience_id")
            if not exp_id:
                continue
            if exp_id not in experiences:
                experiences[exp_id] = {
                    "id": exp_id,
                    "company": meta.get("company", ""),
                    "title": meta.get("title", ""),
                    "date_range": meta.get("date_range", ""),
                    "summary": "",
                    "bullets": [],
                    "skills": meta.get("skills", []),
                }
            chunk_type = meta.get("chunk_type")
            if chunk_type == "summary":
                experiences[exp_id]["summary"] = doc.get("text", "")
            elif chunk_type == "bullet":
                experiences[exp_id]["bullets"].append(doc.get("text", ""))

        return {"experiences": list(experiences.values())}

    async def upload_resume(
        self, user_id: UUID, resume: Optional[Resume] = None, resume_text: Optional[str] = None
    ) -> dict:
        """
        Upload and embed a user's resume into the vector store.
        Supports structured Resume object or plain text.
        Returns the number of documents embedded.
        """
        if resume:
            self.memory_store.clear_experiences(str(user_id))
            count = await self.embedder.embed_resume(str(user_id), resume)
            resume_id = db.save_resume(
                user_id=str(user_id),
                source_type="structured",
                parsed=resume.model_dump(mode="json"),
                embedded_count=count,
            )
        elif resume_text:
            # Plain text mode: split into chunks and store directly
            count = await self._embed_plain_text(str(user_id), resume_text)
            resume_id = db.save_resume(
                user_id=str(user_id),
                source_type="text",
                raw_text=resume_text,
                parsed=self._parse_standard_resume_text(resume_text),
                embedded_count=count,
            )
        else:
            return {"success": False, "resume_id": None, "embedded_count": 0, "message": "No resume content provided."}

        return {
            "success": True,
            "resume_id": resume_id,
            "embedded_count": count,
            "message": f"Resume uploaded and {count} chunks embedded.",
        }

    async def upload_resume_file(self, user_id: UUID, filename: str, file_bytes: bytes) -> dict:
        """
        Upload a resume file (.docx / .pdf / .txt), parse to text, and embed.
        """
        try:
            text = parse_resume_file(filename, file_bytes)
        except Exception as e:
            return {"success": False, "embedded_count": 0, "message": f"Failed to parse file: {e}"}

        count = await self._embed_plain_text(str(user_id), text)
        resume_id = db.save_resume(
            user_id=str(user_id),
            source_type="file",
            filename=filename,
            raw_text=text,
            parsed=self._parse_standard_resume_text(text),
            embedded_count=count,
        )
        return {
            "success": True,
            "resume_id": resume_id,
            "embedded_count": count,
            "message": f"File '{filename}' parsed and {count} chunks embedded.",
        }

    def get_latest_resume(self, user_id: UUID) -> dict | None:
        return db.get_latest_resume(str(user_id))

    async def _embed_plain_text(self, user_id: str, text: str) -> int:
        """Split plain text into chunks and store in Chroma."""
        self.memory_store.clear_experiences(user_id)
        # Simple chunking: split by blank lines, then by sentences if too long
        raw_chunks = [chunk.strip() for chunk in text.split("\n\n") if chunk.strip()]
        documents = [{
            "text": text,
            "metadata": {"chunk_type": "raw_resume", "source": "user_upload"},
        }]
        for idx, chunk in enumerate(raw_chunks):
            # If chunk is too long (>500 chars), split by sentences
            if len(chunk) > 500:
                sentences = chunk.replace(". ", ".\n").split("\n")
                for s_idx, sentence in enumerate(sentences):
                    s = sentence.strip()
                    if s:
                        documents.append({
                            "text": s,
                            "metadata": {
                                "chunk_type": "text",
                                "paragraph_index": idx,
                                "sentence_index": s_idx,
                                "source": "user_upload",
                            },
                        })
            else:
                documents.append({
                    "text": chunk,
                    "metadata": {
                        "chunk_type": "text",
                        "paragraph_index": idx,
                        "source": "user_upload",
                    },
                })

        if not documents:
            return 0

        await self.embedder.store.add_experiences(user_id, documents)
        return len(documents)

    def _parse_standard_resume_text(self, text: str) -> dict:
        """Parse the user's standard one-page resume text into structured sections."""
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = normalized.replace("\u2022 COMPETITIONS", "COMPETITIONS")
        normalized = normalized.replace("• COMPETITIONS", "COMPETITIONS")
        lines = [line.strip() for line in normalized.splitlines() if line.strip()]
        if not lines:
            return {}

        section_names = [
            "EDUCATION",
            "PROFESSIONAL EXPERIENCE",
            "PROJECTS",
            "COMPETITIONS",
            "SKILLS & CERTIFICATIONS",
        ]
        section_indexes: dict[str, int] = {}
        for idx, line in enumerate(lines):
            cleaned = line.strip("•* ").upper()
            if cleaned in section_names and cleaned not in section_indexes:
                section_indexes[cleaned] = idx

        first_section_idx = min(section_indexes.values()) if section_indexes else min(len(lines), 3)
        header_lines = lines[:first_section_idx]
        candidate_name = header_lines[0] if header_lines else ""
        contact_line = header_lines[1] if len(header_lines) > 1 else ""
        summary = " ".join(header_lines[2:]) if len(header_lines) > 2 else ""

        def section_lines(name: str) -> list[str]:
            if name not in section_indexes:
                return []
            start = section_indexes[name] + 1
            later = [pos for sec, pos in section_indexes.items() if pos > section_indexes[name]]
            end = min(later) if later else len(lines)
            return lines[start:end]

        return {
            "candidate_name": candidate_name,
            "contact_line": contact_line,
            "summary": summary,
            "education": self._parse_education(section_lines("EDUCATION")),
            "experiences": self._parse_role_entries(section_lines("PROFESSIONAL EXPERIENCE"), "experience"),
            "projects": self._parse_role_entries(section_lines("PROJECTS"), "project"),
            "competitions": self._parse_role_entries(section_lines("COMPETITIONS"), "competition"),
            "skills": self._split_skills(" ".join(section_lines("SKILLS & CERTIFICATIONS"))),
            "skills_certifications": " ".join(section_lines("SKILLS & CERTIFICATIONS")),
        }

    def _parse_education(self, lines: list[str]) -> list[dict]:
        entries = []
        current: dict | None = None
        date_pattern = re.compile(r"(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*-\s*(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}")
        for line in lines:
            if line.startswith("•") or line.startswith("*"):
                if current is not None:
                    coursework = line.strip("•* ")
                    coursework = coursework.replace("Coursework:", "").strip()
                    current["coursework"] = [item.strip() for item in coursework.split("|") if item.strip()]
                continue
            match = date_pattern.search(line)
            if match:
                if current:
                    entries.append(current)
                current = {
                    "institution": line[: match.start()].strip(),
                    "date_range": match.group(0),
                    "degree": "",
                    "field": "",
                    "location": "",
                    "coursework": [],
                }
            elif current is not None:
                current["degree"] = line
        if current:
            entries.append(current)
        return entries

    def _parse_role_entries(self, lines: list[str], entry_type: str) -> list[dict]:
        entries = []
        current: dict | None = None
        for line in lines:
            is_bullet = line.startswith("•") or line.startswith("*")
            if is_bullet:
                if current is not None:
                    current.setdefault("bullets", []).append({"text": line.strip("•* ")})
                continue
            if current is not None and "|" not in line:
                bullets = current.get("bullets") or []
                if bullets:
                    bullets[-1]["text"] = f"{bullets[-1].get('text', '')} {line}".strip()
                    continue
                if entry_type == "project" and line.lower() == "project":
                    current["context"] = f"{current.get('context', '').strip()} Project".strip()
                    continue
            if current:
                entries.append(current)
            if entry_type == "experience":
                title, company, location, date_range = self._parse_heading(line)
                current = {"title": title, "company": company, "location": location, "date_range": date_range, "bullets": []}
            elif entry_type == "project":
                name, tools, context, date_range = self._parse_project_heading(line)
                current = {"name": name, "tools": self._split_skills(tools), "context": context or "Independent Project", "date_range": date_range, "bullets": [], "skills": self._split_skills(tools)}
            else:
                name, role, location, date_range = self._parse_heading(line)
                current = {"name": name, "role": role, "location": location, "date_range": date_range, "bullets": []}
        if current:
            entries.append(current)
        return entries

    def _parse_project_heading(self, line: str) -> tuple[str, str, str, str]:
        parts = [part.strip() for part in line.split("|")]
        if len(parts) >= 3:
            return self._parse_heading(line)
        if len(parts) == 2:
            name = parts[0]
            rest = parts[1]
            context = "Independent Project" if "independent project" in rest.lower() else ""
            tools = re.sub(r"\bIndependent\s+Project\b", "", rest, flags=re.IGNORECASE).strip(" ,")
            return name, tools, context, ""
        return line, "", "Independent Project", ""

    def _parse_heading(self, line: str) -> tuple[str, str, str, str]:
        parts = [part.strip() for part in line.split("|")]
        left = parts[0] if parts else line
        middle = parts[1] if len(parts) > 1 else ""
        tail = parts[2] if len(parts) > 2 else ""
        if not tail and len(parts) == 2:
            tail = middle
            middle = ""
        tail_parts = [part.strip() for part in re.split(r"\s+-\s+|\s+—\s+", tail) if part.strip()]
        if len(tail_parts) >= 2:
            return left, middle, " - ".join(tail_parts[:-1]), tail_parts[-1]
        return left, middle, tail, ""

    def _split_skills(self, value: str) -> list[str]:
        return [item.strip(" .") for item in re.split(r"[,;]", value or "") if item.strip(" .")]

    async def tailor(
        self,
        user_id: UUID,
        resume_id: UUID,
        jd_text: str,
        job_id: Optional[UUID] = None,
    ) -> dict:
        """
        Run the full tailoring pipeline.
        """
        # Load resume data from vector store
        resume_data = self._rebuild_resume_data(str(user_id))

        # Prepare initial state for LangGraph
        initial_state = {
            "user_id": str(user_id),
            "resume_id": str(resume_id),
            "user_input": jd_text,
            "resume_data": resume_data,
            "jd_text": jd_text,
            "jd_parsed": None,
            "matched_experiences": None,
            "tailored_resume": None,
            "evidence_check": None,
            "agent_response": "",
            "requires_clarification": False,
            "memory_context": {},
        }

        # Execute the agent graph
        result = await tailor_agent.ainvoke(initial_state)

        tailored_resume = result.get("tailored_resume") or {}
        jd_parsed = result.get("jd_parsed") or {}
        markdown = self.text_exporter.render(tailored_resume)
        key_map = self._build_key_map(jd_parsed, tailored_resume, result.get("matched_experiences") or [])
        draft = draft_store.create(
            user_id=str(user_id),
            resume_id=str(resume_id),
            jd_text=jd_text,
            jd_parsed=jd_parsed,
            tailored_resume=tailored_resume,
            markdown=markdown,
            key_map=key_map,
        )

        tailored_resume_id = db.save_tailored_resume(
            user_id=str(user_id),
            resume_id=str(resume_id),
            job_id=str(job_id) if job_id else None,
            jd_text=jd_text,
            jd_parsed=jd_parsed,
            tailored_resume=tailored_resume,
            markdown=markdown,
            key_map=key_map,
        )
        draft["tailored_resume_id"] = tailored_resume_id
        db.save_draft(draft, tailored_resume_id=tailored_resume_id)
        await event_bus.publish(
            ResumeTailoredEvent(
                user_id=user_id,
                resume_id=resume_id,
                tailored_resume_id=UUID(tailored_resume_id),
                job_id=job_id,
                tailoring_summary={"message": tailored_resume.get("tailoring_summary", "")},
                ats_score_estimate=tailored_resume.get("ats_score_estimate"),
            )
        )

        return {
            "success": not result.get("requires_clarification", False),
            "tailored_resume": tailored_resume,
            "message": result.get("agent_response", ""),
            "clarification_needed": result.get("requires_clarification", False),
            "clarification_question": result.get("agent_response") if result.get("requires_clarification") else None,
            "ats_score_estimate": tailored_resume.get("ats_score_estimate") if tailored_resume else None,
            "tailored_resume_id": tailored_resume_id,
            "draft_id": draft["draft_id"],
            "revision_id": draft["current_revision_id"],
            "markdown": markdown,
            "key_map": key_map,
        }

    async def parse_jd(self, jd_text: str) -> ParsedJobDescription:
        """
        Standalone JD parsing utility.
        """
        return await self.jd_parser.parse(jd_text)

    def export_text(self, tailored_resume: dict) -> str:
        """
        Export tailored resume as plain text.
        """
        return self.text_exporter.render(tailored_resume)

    async def modify_draft(self, user_id: UUID, draft_id: str, instruction: str) -> dict:
        draft = draft_store.get(draft_id)
        if not draft:
            return {"success": False, "draft_id": draft_id, "message": "Draft not found."}
        if draft.get("user_id") != str(user_id):
            return {"success": False, "draft_id": draft_id, "message": "Draft does not belong to user."}

        original_resume = self._rebuild_resume_data(str(user_id))
        revised = await self.tailor_node.revise(
            current_resume=draft.get("tailored_resume") or {},
            instruction=instruction,
            jd_parsed=draft.get("jd_parsed") or {},
            key_map=draft.get("key_map") or [],
            original_resume=original_resume,
        )
        markdown = self.text_exporter.render(revised)
        key_map = self._build_key_map(draft.get("jd_parsed") or {}, revised, [])
        updated = draft_store.update(
            draft_id=draft_id,
            instruction=instruction,
            tailored_resume=revised,
            markdown=markdown,
            key_map=key_map,
        )
        db.save_draft(updated, tailored_resume_id=updated.get("tailored_resume_id"))
        return {
            "success": True,
            "draft_id": draft_id,
            "revision_id": updated["current_revision_id"],
            "tailored_resume": revised,
            "markdown": markdown,
            "key_map": key_map,
            "message": "Updated the resume draft based on your instruction.",
        }

    def get_draft(self, user_id: UUID, draft_id: str) -> dict | None:
        draft = draft_store.get(draft_id)
        if not draft:
            draft = db.get_draft(draft_id, user_id=str(user_id))
        if not draft or draft.get("user_id") != str(user_id):
            return None
        return draft

    def get_tailored_resume(self, tailored_resume_id: UUID, user_id: UUID | None = None) -> dict | None:
        return db.get_tailored_resume(str(tailored_resume_id), str(user_id) if user_id else None)

    def _build_key_map(self, jd_parsed: dict, tailored_resume: dict, matched_experiences: list) -> list[dict]:
        keys: list[str] = []
        for field in ["required_skills", "preferred_skills", "ats_keywords"]:
            for item in jd_parsed.get(field, []) or []:
                if isinstance(item, str) and item.strip():
                    keys.append(item.strip())
        for item in jd_parsed.get("key_responsibilities", []) or []:
            if isinstance(item, str) and item.strip():
                keys.append(item.strip())

        deduped = []
        seen = set()
        for key in keys:
            lower = key.lower()
            if lower not in seen:
                deduped.append(key)
                seen.add(lower)

        resume_text = self.text_exporter.render(tailored_resume)
        resume_lower = resume_text.lower()
        results = []
        for key in deduped[:24]:
            key_lower = key.lower()
            tokens = [t for t in key_lower.replace("/", " ").replace("-", " ").split() if len(t) > 2]
            matched_terms = []
            if key_lower in resume_lower:
                matched_terms.append(key)
            else:
                matched_terms.extend([t for t in tokens if t in resume_lower])

            resume_phrase = self._find_resume_phrase(resume_text, matched_terms)
            status = "matched" if key_lower in resume_lower else "partial" if matched_terms else "missing"
            results.append({
                "jd_key": key,
                "resume_phrase": resume_phrase,
                "status": status,
                "highlight_terms": matched_terms[:6],
                "note": self._key_map_note(status),
            })
        return results

    def _find_resume_phrase(self, resume_text: str, terms: list[str]) -> str:
        if not terms:
            return "No supported resume phrase found yet."
        normalized_terms = [t.lower() for t in terms]
        for raw_line in resume_text.splitlines():
            line = raw_line.strip(" •")
            lower = line.lower()
            if line and any(term.lower() in lower for term in normalized_terms):
                return line[:220]
        return ", ".join(terms)

    def _key_map_note(self, status: str) -> str:
        if status == "matched":
            return "Directly reflected in the tailored resume."
        if status == "partial":
            return "Covered through adjacent, evidence-backed wording."
        return "Not added because the original resume does not support it."

    def export_draft_docx(self, draft: dict) -> bytes:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError as e:
            raise RuntimeError("python-docx is required for Word export") from e

        document = Document()
        tr = draft.get("tailored_resume") or {}
        for section in document.sections:
            section.top_margin = Inches(0.45)
            section.bottom_margin = Inches(0.45)
            section.left_margin = Inches(0.55)
            section.right_margin = Inches(0.55)

        if tr.get("candidate_name"):
            paragraph = document.add_paragraph()
            paragraph.alignment = 1
            paragraph.add_run(str(tr["candidate_name"]).upper()).bold = True
        if tr.get("contact_line"):
            paragraph = document.add_paragraph()
            paragraph.alignment = 1
            paragraph.add_run(str(tr["contact_line"]))
        if tr.get("summary"):
            document.add_paragraph(tr["summary"])

        if tr.get("education"):
            document.add_heading("EDUCATION", level=2)
            for edu in tr["education"]:
                header = document.add_paragraph()
                header.add_run(" ".join(p for p in [edu.get("institution", ""), edu.get("date_range", "")] if p)).bold = True
                detail = " ".join(str(p) for p in [edu.get("degree"), edu.get("field"), edu.get("location")] if p)
                if detail:
                    document.add_paragraph(detail)
                coursework = edu.get("coursework") or []
                if coursework:
                    document.add_paragraph(f"Coursework: {' | '.join(str(c) for c in coursework)}", style="List Bullet")

        if tr.get("experiences"):
            document.add_heading("PROFESSIONAL EXPERIENCE", level=2)
            for exp in tr["experiences"]:
                header = document.add_paragraph()
                header.add_run(
                    self._docx_heading(exp.get("title"), exp.get("company"), exp.get("location"), exp.get("date_range"))
                ).bold = True
                for bullet in exp.get("bullets", [])[:3]:
                    text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                    document.add_paragraph(text, style="List Bullet")

        if tr.get("projects"):
            document.add_heading("PROJECTS", level=2)
            for proj in tr["projects"]:
                header = document.add_paragraph()
                tools = ", ".join(str(t) for t in (proj.get("tools") or proj.get("skills") or []))
                header.add_run(self._docx_heading(proj.get("name"), tools, proj.get("context", "Independent Project"), proj.get("date_range"))).bold = True
                bullets = proj.get("bullets") or []
                if bullets:
                    for bullet in bullets[:3]:
                        text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                        document.add_paragraph(text, style="List Bullet")
                elif proj.get("description"):
                    document.add_paragraph(proj["description"], style="List Bullet")

        if tr.get("competitions"):
            document.add_heading("COMPETITIONS", level=2)
            for comp in tr["competitions"]:
                header = document.add_paragraph()
                header.add_run(self._docx_heading(comp.get("name"), comp.get("role"), comp.get("location"), comp.get("date_range"))).bold = True
                for bullet in comp.get("bullets", []):
                    text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                    document.add_paragraph(text, style="List Bullet")

        skills_text = tr.get("skills_certifications") or ", ".join(
            str(s) for s in [*(tr.get("skills") or []), *(tr.get("certifications") or [])]
        )
        if skills_text:
            document.add_heading("SKILLS & CERTIFICATIONS", level=2)
            document.add_paragraph(skills_text)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _docx_heading(self, left: object, middle: object, location: object, date_range: object) -> str:
        first = " | ".join(str(p) for p in [left, middle] if p)
        second = " - ".join(str(p) for p in [location, date_range] if p)
        return " ".join(p for p in [first, second] if p)

    def export_draft_pdf(self, draft: dict) -> bytes:
        markdown = draft.get("markdown") or "No resume content available."
        raw_lines = [line.strip() for line in markdown.splitlines() if line.strip()]
        configs = [
            {"font_size": 7.8, "leading": 9.1, "wrap_width": 132},
            {"font_size": 7.2, "leading": 8.3, "wrap_width": 146},
            {"font_size": 6.7, "leading": 7.7, "wrap_width": 158},
        ]
        page_height = 792
        top = 752
        bottom = 34

        selected_lines: list[str] = []
        selected = configs[-1]
        for config in configs:
            lines = self._wrap_resume_lines(raw_lines, config["wrap_width"])
            max_lines = int((top - bottom) / config["leading"])
            if len(lines) <= max_lines:
                selected_lines = lines
                selected = config
                break
            selected_lines = lines
            selected = config

        return self._simple_pdf_one_page(
            selected_lines,
            font_size=selected["font_size"],
            leading=selected["leading"],
            top=top,
            page_height=page_height,
        )

    def _wrap_resume_lines(self, raw_lines: list[str], width: int) -> list[str]:
        lines: list[str] = []
        section_names = {
            "EDUCATION",
            "PROFESSIONAL EXPERIENCE",
            "PROJECTS",
            "COMPETITIONS",
            "SKILLS & CERTIFICATIONS",
        }
        for raw in raw_lines:
            line = raw.replace("* ", "• ")
            if line in section_names:
                lines.append(line)
                continue
            is_bullet = line.startswith("•")
            wrapped = wrap(line, width=width, subsequent_indent="  " if is_bullet else "") or [line]
            lines.extend(wrapped)
        return lines

    def _simple_pdf_one_page(
        self,
        lines: list[str],
        *,
        font_size: float,
        leading: float,
        top: int,
        page_height: int,
    ) -> bytes:
        objects: list[bytes] = []

        def add(obj: str) -> int:
            objects.append(obj.encode("cp1252", errors="replace"))
            return len(objects)

        add("<< /Type /Catalog /Pages 2 0 R >>")
        add("<< /Type /Pages /Kids [] /Count 0 >>")
        font_id = add("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        page_ids = []
        commands = ["BT", f"/F1 {font_size} Tf", f"48 {top} Td", f"{leading} TL"]
        for line in lines:
            safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            commands.append(f"({safe}) Tj")
            commands.append("T*")
        commands.append("ET")
        stream = "\n".join(commands)
        content_id = add(f"<< /Length {len(stream.encode('cp1252', errors='replace'))} >>\nstream\n{stream}\nendstream")
        page_id = add(f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 {page_height}] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>")
        page_ids.append(page_id)
        objects[1] = f"<< /Type /Pages /Kids [{' '.join(f'{pid} 0 R' for pid in page_ids)}] /Count {len(page_ids)} >>".encode("latin-1")

        pdf = bytearray(b"%PDF-1.4\n")
        offsets = []
        for idx, obj in enumerate(objects, start=1):
            offsets.append(len(pdf))
            pdf.extend(f"{idx} 0 obj\n".encode("latin-1"))
            pdf.extend(obj)
            pdf.extend(b"\nendobj\n")
        xref_at = len(pdf)
        pdf.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("latin-1"))
        for offset in offsets:
            pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
        pdf.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF".encode("latin-1"))
        return bytes(pdf)

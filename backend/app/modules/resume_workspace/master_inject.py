"""Inject tailored content into master DOCX while preserving paragraph/run formatting."""

from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from typing import Any


SECTION_HEADINGS = {
    "EDUCATION",
    "PROFESSIONAL EXPERIENCE",
    "PROJECTS",
    "COMPETITIONS",
    "SKILLS & CERTIFICATIONS",
}


def _set_paragraph_text_keep_format(paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _norm(s: str) -> str:
    return " ".join((s or "").split()).strip().lower()


def _is_section_heading(text: str) -> bool:
    t = text.strip().upper()
    return t in SECTION_HEADINGS


def _is_entry_heading(text: str) -> bool:
    # Title | Company  OR  Name | Tools
    if "|" in text and not text.strip().startswith(("•", "*", "-")):
        return True
    return False


def inject_content(master_docx: bytes, tailored: dict[str, Any], master_inventory: dict[str, Any]) -> bytes:
    from docx import Document

    doc = Document(BytesIO(master_docx))
    tailored = tailored or {}
    inventory = master_inventory or {}

    bullet_map: dict[str, str] = {}
    for section in ("experiences", "projects"):
        inv_entries = {
            _entry_key(e, section): e
            for e in (inventory.get(section) or [])
            if isinstance(e, dict)
        }
        for entry in tailored.get(section) or []:
            if not isinstance(entry, dict):
                continue
            key = _entry_key(entry, section)
            inv = inv_entries.get(key)
            t_bullets = _bullets(entry)
            i_bullets = _bullets(inv) if inv else []
            for i, tb in enumerate(t_bullets):
                new_t = str(tb.get("text") or "")
                if not new_t:
                    continue
                orig = str(tb.get("original_text") or "")
                if orig:
                    bullet_map[_norm(orig)] = new_t
                if i < len(i_bullets):
                    inv_t = str(i_bullets[i].get("text") or "")
                    if inv_t:
                        bullet_map[_norm(inv_t)] = new_t

    new_summary = str(tailored.get("summary") or "").strip()
    new_skills = str(tailored.get("skills_certifications") or "").strip()
    old_summary = str(inventory.get("summary") or "").strip()
    old_skills = str(inventory.get("skills_certifications") or "").strip()

    # Track which section we're in while scanning
    current_section = ""
    for p in doc.paragraphs:
        text = p.text
        raw = text.strip()
        if not raw:
            continue

        if _is_section_heading(raw):
            current_section = raw.upper()
            continue

        # Never touch headings / entry headers
        if _is_entry_heading(raw):
            continue

        n = _norm(text)

        # Summary: only before EDUCATION, long paragraph matching inventory summary
        if (
            new_summary
            and current_section == ""
            and old_summary
            and len(raw) > 80
            and (
                n == _norm(old_summary)
                or _norm(old_summary)[:50] in n
                or n.startswith("data science m.s.")
                or n.startswith("data analyst candidate")
            )
        ):
            _set_paragraph_text_keep_format(p, new_summary)
            continue

        # Skills: only inside SKILLS section, or exact inventory skills match
        if new_skills and (
            current_section.startswith("SKILLS")
            or (old_skills and n == _norm(old_skills))
        ):
            if raw.count(",") >= 3 and not _is_entry_heading(raw):
                _set_paragraph_text_keep_format(p, new_skills)
                continue

        # Bullets only
        if current_section in {"PROFESSIONAL EXPERIENCE", "PROJECTS", "COMPETITIONS", "EDUCATION"}:
            stripped = raw.lstrip("•*-–— ").strip()
            ns = _norm(stripped)
            if ns in bullet_map:
                # Preserve list style; don't prepend glyph if style already bullets
                _set_paragraph_text_keep_format(p, bullet_map[ns])
                continue
            for old_n, new_t in bullet_map.items():
                if len(old_n) > 60 and (old_n[:70] in ns or ns[:70] in old_n):
                    _set_paragraph_text_keep_format(p, new_t)
                    break

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def content_integrity_check(docx_bytes: bytes, inventory: dict[str, Any]) -> dict[str, Any]:
    """Ensure project/experience titles were not overwritten by skills dumps."""
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    text = text.replace("\u00a0", " ").replace("\u2009", " ")
    errors: list[str] = []

    for exp in inventory.get("experiences") or []:
        company = str(exp.get("company") or "").replace("\u00a0", " ")
        if company and company not in text:
            # fuzzy: first significant token
            token = company.split()[0] if company.split() else ""
            if token and token not in text:
                errors.append(f"missing_experience_company:{company}")

    for proj in inventory.get("projects") or []:
        name = str(proj.get("name") or "")
        if name and name not in text:
            errors.append(f"missing_project_title:{name}")

    # Skills dump should not appear as a project heading line with pipes removed oddly
    skills = str(inventory.get("skills_certifications") or "")
    if skills:
        # If a paragraph equals full skills AND sits under projects as title-like — detected if
        # skills string appears more than twice (skills section + accidental dupes)
        count = text.count(skills[:40]) if len(skills) > 40 else text.count(skills)
        if count > 2:
            errors.append("skills_string_duplicated_excessively")

    return {"ok": len(errors) == 0, "errors": errors}


def _entry_key(entry: dict[str, Any], section: str) -> str:
    if section == "experiences":
        return f"{entry.get('company','')}|{entry.get('title','')}"
    return str(entry.get("name") or "")


def _bullets(entry: dict[str, Any] | None) -> list[dict[str, Any]]:
    if not entry:
        return []
    out = []
    for b in entry.get("bullets") or []:
        if isinstance(b, dict):
            out.append(b)
        else:
            out.append({"text": str(b)})
    return out


def clone_inventory(inventory: dict[str, Any]) -> dict[str, Any]:
    return deepcopy(inventory)

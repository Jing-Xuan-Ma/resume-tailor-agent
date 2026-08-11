import re
import tempfile
from io import BytesIO
from pathlib import Path


class ResumeTemplateEditor:
    def __init__(self):
        self._run_map: dict[str, tuple[int, int]] = {}
        self._paragraph_count = 0

    def load_template(self, docx_bytes: bytes) -> dict:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for template editing") from exc

        doc = Document(BytesIO(docx_bytes))
        run_map = {}
        editable_blocks = []

        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                block_id = f"p{p_idx}_r{r_idx}"
                run_map[block_id] = (p_idx, r_idx)
                if run.text.strip():
                    font_name = run.font.name
                    font_size = None
                    if run.font.size:
                        font_size = str(int(run.font.size.pt))
                    editable_blocks.append(
                        {
                            "block_id": block_id,
                            "paragraph_index": p_idx,
                            "run_index": r_idx,
                            "text": run.text,
                            "bold": bool(run.bold),
                            "font_name": font_name,
                            "font_size": font_size,
                        }
                    )

        self._run_map = run_map
        self._paragraph_count = len(doc.paragraphs)
        return {
            "blocks": editable_blocks,
            "paragraph_count": self._paragraph_count,
            "block_count": len(editable_blocks),
        }

    def apply_text_replacements(self, docx_bytes: bytes, replacements: dict[str, str]) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for template editing") from exc

        doc = Document(BytesIO(docx_bytes))

        for block_id, new_text in replacements.items():
            if block_id not in self._run_map:
                continue
            p_idx, r_idx = self._run_map[block_id]
            paragraph = doc.paragraphs[p_idx]
            if r_idx < len(paragraph.runs):
                paragraph.runs[r_idx].text = new_text

        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def _run_id_for_paragraph_text(self, paragraph_text: str, template_docx: bytes) -> str | None:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required") from exc

        doc = Document(BytesIO(template_docx))
        # Rebuild run map if not already loaded
        if not self._run_map:
            self.load_template(template_docx)

        needle = paragraph_text.strip().lower()
        for block_id, (p_idx, r_idx) in self._run_map.items():
            paragraph = doc.paragraphs[p_idx]
            if r_idx < len(paragraph.runs):
                run_text = paragraph.runs[r_idx].text.strip().lower()
                if run_text == needle or (len(needle) > 20 and needle in run_text):
                    return block_id
        return None

    def build_replacement_map(
        self, original_resume: dict, tailored_resume: dict, template_docx: bytes
    ) -> dict[str, str]:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required") from exc

        doc = Document(BytesIO(template_docx))
        if not self._run_map:
            self.load_template(template_docx)

        replacements: dict[str, str] = {}
        text_blocks = self._extract_text_blocks(tailored_resume)

        for block_id, (p_idx, r_idx) in self._run_map.items():
            paragraph = doc.paragraphs[p_idx]
            if r_idx >= len(paragraph.runs):
                continue
            run = paragraph.runs[r_idx]
            original_text = run.text.strip()
            if not original_text:
                continue

            new_text = self._find_best_match(original_text, text_blocks)
            if new_text and new_text != original_text:
                replacements[block_id] = new_text

        return replacements

    def _extract_text_blocks(self, resume: dict) -> dict[str, str]:
        blocks: dict[str, str] = {}
        if resume.get("summary"):
            blocks["summary"] = resume["summary"]
        for exp in resume.get("experiences", []):
            title = exp.get("title", "")
            company = exp.get("company", "")
            heading = f"{title} | {company}" if title and company else (title or company)
            if heading:
                blocks[f"exp_heading_{company}"] = heading
            for i, bullet in enumerate(exp.get("bullets", [])):
                text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                if text:
                    blocks[f"exp_bullet_{company}_{i}"] = text
        for proj in resume.get("projects", []):
            name = proj.get("name", "")
            tools = ", ".join(str(t) for t in (proj.get("tools") or []))
            heading = f"{name} | {tools}" if name and tools else name
            if heading:
                blocks[f"proj_heading_{name}"] = heading
            for i, bullet in enumerate(proj.get("bullets", [])):
                text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                if text:
                    blocks[f"proj_bullet_{name}_{i}"] = text
        if resume.get("skills_certifications"):
            blocks["skills"] = resume["skills_certifications"]
        return blocks

    def _find_best_match(self, original_text: str, text_blocks: dict[str, str]) -> str | None:
        original_lower = original_text.lower()
        best_match = None
        best_score = 0

        for _key, new_text in text_blocks.items():
            score = self._similarity_score(original_lower, new_text.lower())
            if score > best_score:
                best_score = score
                best_match = new_text

        return best_match

    def _similarity_score(self, a: str, b: str) -> float:
        if not a or not b:
            return 0
        if a == b:
            return 1.0
        a_words = set(a.split())
        b_words = set(b.split())
        if not a_words or not b_words:
            return 0
        intersection = a_words & b_words
        union = a_words | b_words
        return len(intersection) / len(union)

    @staticmethod
    def generate_preview_pdf(full_resume: dict) -> bytes:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError as exc:
            raise RuntimeError("python-docx is required") from exc

        document = Document()

        for section in document.sections:
            section.top_margin = Inches(0.2)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0.2)
            section.right_margin = Inches(0.2)

        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        fr = full_resume

        if fr.get("candidate_name"):
            p = document.add_paragraph()
            p.alignment = 1
            run = p.add_run(str(fr["candidate_name"]).upper())
            run.bold = True
            run.font.size = Pt(15)
            run.font.name = "Calibri"

        if fr.get("contact_line"):
            p = document.add_paragraph()
            p.alignment = 1
            run = p.add_run(str(fr["contact_line"]))
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        if fr.get("summary"):
            p = document.add_paragraph()
            p.alignment = 3
            run = p.add_run(str(fr["summary"]))
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        section_defs = [
            ("education", "EDUCATION", "institution"),
            ("experiences", "PROFESSIONAL EXPERIENCE", "company"),
            ("projects", "PROJECTS", "name"),
        ]

        for key, title, _ in section_defs:
            items = fr.get(key, [])
            if not items:
                continue
            document.add_paragraph()
            p = document.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(0)

            for item in items:
                if key == "education":
                    edu_bits = [item.get("institution", ""), item.get("date_range", "")]
                    heading = " | ".join(str(p) for p in edu_bits if p)
                    if heading:
                        p = document.add_paragraph()
                        run = p.add_run(heading)
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"
                    detail = " ".join(str(p) for p in [item.get("degree"), item.get("field")] if p)
                    if detail:
                        p = document.add_paragraph()
                        run = p.add_run(detail)
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"
                        p.alignment = 3
                else:
                    left = item.get("title") or item.get("name", "")
                    middle = item.get("company") or ", ".join(
                        str(t) for t in (item.get("tools") or [])
                    )
                    heading = " | ".join(str(p) for p in [left, middle] if p)
                    if item.get("date_range"):
                        heading += f" - {item['date_range']}"
                    if heading:
                        p = document.add_paragraph()
                        run = p.add_run(heading)
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"
                        p.paragraph_format.space_after = Pt(0)

                    for bullet in item.get("bullets", []):
                        text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                        if text:
                            p = document.add_paragraph()
                            run = p.add_run(f"• {text}")
                            run.font.size = Pt(10)
                            run.font.name = "Calibri"
                            p.alignment = 3
                            p.paragraph_format.space_after = Pt(0)

        skills = fr.get("skills_certifications")
        if skills:
            document.add_paragraph()
            p = document.add_paragraph()
            run = p.add_run("SKILLS & CERTIFICATIONS")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

            p = document.add_paragraph()
            run = p.add_run(str(skills))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            p.alignment = 3

        docx_bytes = BytesIO()
        document.save(docx_bytes)
        return docx_bytes.getvalue()

    @staticmethod
    def generate_export_docx(full_resume: dict, template_docx: bytes | None = None) -> bytes:
        if template_docx:
            return template_docx
        return ResumeTemplateEditor.generate_preview_pdf(full_resume)

    @staticmethod
    def convert_docx_to_pdf_via_word(docx_bytes: bytes, *, label: str = "preview") -> bytes:
        """True master-layout PDF via Microsoft Word COM (Windows)."""
        import sys
        from pathlib import Path

        repo_root = Path(__file__).resolve().parents[4]
        scripts_dir = repo_root / "scripts"
        if str(scripts_dir) not in sys.path:
            sys.path.insert(0, str(scripts_dir))
        from rg_word_pdf import word_export_pdf  # type: ignore

        with tempfile.TemporaryDirectory(prefix="ws_word_pdf_") as td:
            pdf_out = Path(td) / f"{label}.pdf"
            word_export_pdf(docx_bytes, pdf_out, label=label)
            return pdf_out.read_bytes()

    @staticmethod
    def _find_soffice_binary() -> str:
        import shutil

        on_path = shutil.which("soffice") or shutil.which("libreoffice")
        if on_path:
            return on_path
        candidates = [
            Path.home() / "Applications" / "LibreOffice.app" / "Contents" / "MacOS" / "soffice",
            Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
            Path("/usr/bin/soffice"),
            Path("/opt/homebrew/bin/soffice"),
        ]
        for c in candidates:
            if c.exists():
                return str(c)
        return "soffice"

    @staticmethod
    def convert_to_pdf_via_libreoffice(docx_bytes: bytes, timeout_seconds: int = 60) -> bytes:
        import subprocess
        from pathlib import Path

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_docx.write(docx_bytes)
            docx_path = tmp_docx.name

        pdf_path = docx_path.replace(".docx", ".pdf")
        soffice_bin = ResumeTemplateEditor._find_soffice_binary()

        try:
            result = subprocess.run(
                [
                    soffice_bin,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(Path(pdf_path).parent),
                    docx_path,
                ],
                capture_output=True,
                timeout=timeout_seconds,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.decode()}")
            with open(pdf_path, "rb") as f:
                return f.read()
        except FileNotFoundError as exc:
            raise RuntimeError(
                "LibreOffice not found. Install it and ensure 'soffice' is on PATH."
            ) from exc
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError(
                f"LibreOffice conversion timed out after {timeout_seconds}s"
            ) from exc
        finally:
            Path(docx_path).unlink(missing_ok=True)
            Path(pdf_path).unlink(missing_ok=True)

    @staticmethod
    def generate_pdf_from_resume(full_resume: dict) -> bytes:
        """Plain-text PDF fallback — no Markdown markers (##, **, #).

        Does not mutate resume structure. Prefer Word master + enforce_one_page
        for delivery PDFs; this path only shrinks type / hard-clips overflow.
        """
        raw_lines = ResumeTemplateEditor._render_plain_lines(full_resume)

        configs = [
            {"font_size": 9.0, "leading": 11.0, "wrap_width": 92},
            {"font_size": 8.2, "leading": 10.0, "wrap_width": 100},
            {"font_size": 7.4, "leading": 9.0, "wrap_width": 110},
        ]
        page_height = 792
        top = 752
        bottom = 36

        selected_lines = None
        selected = configs[-1]
        for config in configs:
            wrapped = ResumeTemplateEditor._wrap_lines(raw_lines, config["wrap_width"])
            max_lines = int((top - bottom) / config["leading"])
            if len(wrapped) <= max_lines:
                selected_lines = wrapped
                selected = config
                break
        if selected_lines is None:
            wrapped = ResumeTemplateEditor._wrap_lines(raw_lines, selected["wrap_width"])
            max_lines = int((top - bottom) / selected["leading"])
            selected_lines = wrapped[:max_lines]

        return ResumeTemplateEditor._render_pdf(
            selected_lines,
            font_size=selected["font_size"],
            leading=selected["leading"],
            top=top,
            page_height=page_height,
            bottom=bottom,
        )

    @staticmethod
    def _strip_md(text: str) -> str:
        """Remove Markdown emphasis / heading markers from display text."""
        s = str(text or "")
        s = re.sub(r"^#{1,6}\s*", "", s)
        s = s.replace("**", "").replace("__", "").replace("`", "")
        # Constitution: no ATS-hostile bullets / arrows
        for ch in ("→", "←", "⇒", "⇐", "➜", "➔", "➡", "●", "◆", "■", "★", "✓", "✔", "✗"):
            s = s.replace(ch, "-")
        return s.strip()

    @staticmethod
    def _render_plain_lines(resume: dict) -> list[str]:
        """Clean resume lines for PDF preview (no Markdown syntax)."""
        lines: list[str] = []
        if resume.get("candidate_name"):
            lines.append(str(resume["candidate_name"]).upper())
        if resume.get("contact_line"):
            lines.append(ResumeTemplateEditor._strip_md(str(resume["contact_line"])))
        if resume.get("summary"):
            lines.append("")
            lines.append(ResumeTemplateEditor._strip_md(str(resume["summary"])))
        if resume.get("education"):
            lines.append("")
            lines.append("EDUCATION")
            for edu in resume["education"]:
                parts = [p for p in [edu.get("institution"), edu.get("date_range")] if p]
                if parts:
                    lines.append(ResumeTemplateEditor._strip_md(" | ".join(str(p) for p in parts)))
                degree_bits = [
                    p for p in [edu.get("degree"), edu.get("field"), edu.get("location")] if p
                ]
                if degree_bits:
                    lines.append(
                        ResumeTemplateEditor._strip_md(" | ".join(str(p) for p in degree_bits))
                    )
        for key, title in [
            ("experiences", "PROFESSIONAL EXPERIENCE"),
            ("projects", "PROJECTS"),
            ("competitions", "COMPETITIONS"),
        ]:
            items = resume.get(key, [])
            if not items:
                continue
            lines.append("")
            lines.append(title)
            for item in items:
                left = item.get("title") or item.get("name", "")
                mid = item.get("company") or ", ".join(str(t) for t in (item.get("tools") or []))
                heading = " | ".join(str(p) for p in [left, mid] if p)
                right_bits = [p for p in [item.get("location"), item.get("date_range")] if p]
                if right_bits:
                    heading = (
                        f"{heading}  —  {' | '.join(str(p) for p in right_bits)}"
                        if heading
                        else " | ".join(str(p) for p in right_bits)
                    )
                if heading:
                    lines.append(ResumeTemplateEditor._strip_md(heading))
                for bullet in item.get("bullets", []):
                    text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
                    text = ResumeTemplateEditor._strip_md(text)
                    if text:
                        lines.append(f"- {text}")
        skills = resume.get("skills_certifications") or ", ".join(resume.get("skills") or [])
        if skills:
            lines.append("")
            lines.append("SKILLS & CERTIFICATIONS")
            lines.append(ResumeTemplateEditor._strip_md(str(skills)))
        return lines

    @staticmethod
    def _render_markdown(resume: dict) -> str:
        # Kept for .txt export / legacy callers — plain text, no MD markers.
        return "\n".join(ResumeTemplateEditor._render_plain_lines(resume))

    @staticmethod
    def _wrap_lines(raw_lines: list[str], width: int) -> list[str]:
        from textwrap import wrap

        section_names = {
            "EDUCATION",
            "PROFESSIONAL EXPERIENCE",
            "PROJECTS",
            "COMPETITIONS",
            "SKILLS & CERTIFICATIONS",
        }
        lines = []
        for raw in raw_lines:
            line = ResumeTemplateEditor._strip_md(raw)
            line = line.replace("* ", "- ")
            if not line:
                lines.append("")
                continue
            if line in section_names or (line.isupper() and len(line) < 40 and " | " not in line):
                lines.append(line)
                continue
            is_bullet = line.startswith("- ") or line.startswith("•")
            wrapped = wrap(line, width=width, subsequent_indent="  " if is_bullet else "") or [line]
            lines.extend(wrapped)
        return lines

    @staticmethod
    def _render_pdf(
        lines: list[str],
        font_size: float,
        leading: float,
        top: int,
        page_height: int,
        bottom: int = 36,
    ) -> bytes:
        objects: list[bytes] = []

        def add(obj: str) -> int:
            objects.append(obj.encode("cp1252", errors="replace"))
            return len(objects)

        section_names = {
            "EDUCATION",
            "PROFESSIONAL EXPERIENCE",
            "PROJECTS",
            "COMPETITIONS",
            "SKILLS & CERTIFICATIONS",
        }

        max_lines = max(1, int((top - bottom) / leading))
        lines = list(lines)[:max_lines]

        add("<< /Type /Catalog /Pages 2 0 R >>")
        add("<< /Type /Pages /Kids [] /Count 0 >>")
        font_reg = add("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        font_bold = add("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")
        commands = ["BT", f"/F1 {font_size} Tf", f"48 {top} Td", f"{leading} TL"]
        current_bold = False
        for i, line in enumerate(lines):
            is_name = i == 0 and bool(line) and " | " not in line and "@" not in line
            is_section = line in section_names
            want_bold = is_name or is_section
            if want_bold != current_bold:
                commands.append(
                    f"/F{'2' if want_bold else '1'} {font_size + (1.5 if is_name else 0)} Tf"
                )
                current_bold = want_bold
            safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            commands.append(f"({safe}) Tj")
            commands.append("T*")
        commands.append("ET")
        stream = "\n".join(commands)
        stream_len = len(stream.encode("cp1252", errors="replace"))
        content_id = add(f"<< /Length {stream_len} >>\nstream\n{stream}\nendstream")
        page_id = add(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 {page_height}] "
            f"/Resources << /Font << /F1 {font_reg} 0 R /F2 {font_bold} 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        )
        objects[1] = f"<< /Type /Pages /Kids [{page_id} 0 R] /Count 1 >>".encode("latin-1")

        pdf = bytearray(b"%PDF-1.4\n")
        offsets = []
        for idx, obj in enumerate(objects, start=1):
            offsets.append(len(pdf))
            pdf.extend(f"{idx} 0 obj\n".encode("latin-1"))
            pdf.extend(obj)
            pdf.extend(b"\nendobj\n")
        xref_at = len(pdf)
        pdf.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("latin-1"))
        for offset in offsets:
            pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
        trailer = (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF"
        )
        pdf.extend(trailer.encode("latin-1"))
        return bytes(pdf)

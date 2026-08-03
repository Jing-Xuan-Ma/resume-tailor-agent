"""DOCX format fingerprint + comparison for RG format lock."""

from __future__ import annotations

from io import BytesIO
from typing import Any


def fingerprint_docx(docx_bytes: bytes) -> dict[str, Any]:
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    paras = []
    fonts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name if p.style else ""
        paras.append({"style": style, "len": len(text), "empty": not bool(text)})
        for r in p.runs:
            if not r.text.strip():
                continue
            size = int(r.font.size.pt) if r.font.size else None
            fonts.append({"bold": bool(r.bold), "size": size, "name": r.font.name})
            break  # first non-empty run per paragraph

    headings = [p.text.strip().upper() for p in doc.paragraphs if p.text.strip().isupper() and len(p.text.strip()) < 40]
    return {
        "paragraph_count": len(doc.paragraphs),
        "non_empty_paragraphs": sum(1 for p in paras if not p["empty"]),
        "styles": [p["style"] for p in paras],
        "heading_labels": headings,
        "font_sizes": [f["size"] for f in fonts],
        "font_names": [f["name"] for f in fonts],
        "bold_flags": [f["bold"] for f in fonts],
    }


def compare_fingerprints(master: dict[str, Any], generated: dict[str, Any]) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []

    if master["paragraph_count"] != generated["paragraph_count"]:
        errors.append(
            f"paragraph_count {master['paragraph_count']} -> {generated['paragraph_count']}"
        )

    # Required section headings must remain
    required = {"EDUCATION", "PROFESSIONAL EXPERIENCE", "PROJECTS", "SKILLS & CERTIFICATIONS"}
    gen_heads = set(generated.get("heading_labels") or [])
    # COMPETITIONS may be present in master as weird merge; check soft
    missing = required - gen_heads
    # Also accept if heading text appears inside paragraph scan
    if missing:
        # soft: some templates embed COMPETITIONS oddly; required core four
        errors.append(f"missing_headings {sorted(missing)}")

    if master.get("styles") != generated.get("styles"):
        # style sequence drift is hard fail for format lock
        drift = sum(1 for a, b in zip(master["styles"], generated["styles"]) if a != b)
        extra = abs(len(master["styles"]) - len(generated["styles"]))
        if drift or extra:
            errors.append(f"style_sequence_drift drift={drift} extra={extra}")

    # Font size sequence should match for overlapping prefix
    m_sizes = master.get("font_sizes") or []
    g_sizes = generated.get("font_sizes") or []
    for i, (a, b) in enumerate(zip(m_sizes, g_sizes)):
        if a is not None and b is not None and a != b:
            errors.append(f"font_size_mismatch at para-run {i}: {a} -> {b}")
            break

    score = 10
    score -= min(8, len(errors) * 2)
    score -= min(2, len(warnings))
    return {
        "ok": len(errors) == 0,
        "score": max(0, score),
        "errors": errors,
        "warnings": warnings,
    }

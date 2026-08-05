"""JR-1/JR-2 shared job catalog: fingerprint, ingest (write), search (read)."""

from __future__ import annotations

import hashlib
import re
from typing import Any
from urllib.parse import parse_qsl, urlencode, urlparse, urlunparse

from app import db
from app.config import settings
from app.modules.job_discovery.orchestrator import discover_all
from app.modules.job_discovery.quality import assess_listing_quality, filter_quality_leads
from app.modules.job_discovery.scorer import score_job


_UTM_RE = re.compile(r"^utm_", re.I)

# Common public ATS URL shapes → (platform, job_id)
_ATS_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("greenhouse", re.compile(r"(?:boards(?:-api)?\.greenhouse\.io|greenhouse\.io)/([^/]+)/jobs/(\d+)", re.I)),
    ("greenhouse", re.compile(r"job-boards\.greenhouse\.io/([^/]+)/jobs/(\d+)", re.I)),
    ("lever", re.compile(r"jobs\.lever\.co/([^/]+)/([0-9a-f-]{8,})", re.I)),
    ("ashby", re.compile(r"jobs\.ashbyhq\.com/([^/]+)/([0-9a-f-]{8,})", re.I)),
    ("ashby", re.compile(r"api\.ashbyhq\.com/.*?/([0-9a-f-]{8,})", re.I)),
]


def normalize_url(url: str) -> str:
    raw = (url or "").strip()
    if not raw:
        return ""
    parsed = urlparse(raw)
    scheme = (parsed.scheme or "https").lower()
    netloc = parsed.netloc.lower()
    path = parsed.path.rstrip("/") or "/"
    query_pairs = [
        (k, v)
        for k, v in parse_qsl(parsed.query, keep_blank_values=True)
        if not _UTM_RE.match(k)
    ]
    query = urlencode(sorted(query_pairs))
    return urlunparse((scheme, netloc, path, "", query, ""))


def extract_ats_identity(url: str | None) -> tuple[str, str, str] | None:
    """Return (platform, board_or_org, job_id) when URL matches a known ATS."""
    raw = (url or "").strip()
    if not raw:
        return None
    for platform, pattern in _ATS_PATTERNS:
        m = pattern.search(raw)
        if not m:
            continue
        if m.lastindex and m.lastindex >= 2:
            return platform, m.group(1).lower(), m.group(2).lower()
        if m.lastindex == 1:
            return platform, "unknown", m.group(1).lower()
    return None


def listing_fingerprint(
    *,
    source_url: str | None,
    title: str,
    company: str | None,
    source_platform: str | None = None,
) -> str:
    """Stable dedupe key: ATS id → normalized URL → title+company."""
    ats = extract_ats_identity(source_url)
    if ats:
        platform, org, job_id = ats
        payload = f"ats|{platform}|{org}|{job_id}"
        return "ats:" + hashlib.sha1(payload.encode("utf-8")).hexdigest()

    norm = normalize_url(source_url or "")
    if norm:
        return "url:" + hashlib.sha1(norm.encode("utf-8")).hexdigest()

    title_key = re.sub(r"\s+", " ", (title or "").strip().lower())
    company_key = re.sub(r"\s+", " ", (company or "").strip().lower())
    platform_key = (source_platform or "unknown").strip().lower()
    payload = f"{platform_key}|{title_key}|{company_key}"
    return "tc:" + hashlib.sha1(payload.encode("utf-8")).hexdigest()


def infer_work_model(location: str | None, raw_text: str | None = None) -> str:
    blob = f"{location or ''} {raw_text or ''}".lower()
    if re.search(r"\bhybrid\b", blob):
        return "hybrid"
    if re.search(r"\b(on[\s-]?site|onsite|in[\s-]?office)\b", blob):
        return "onsite"
    if re.search(r"\bremote\b", blob) or (location or "").strip().lower() == "remote":
        return "remote"
    return "unknown"


def upsert_lead(lead: dict[str, Any]) -> tuple[str, bool]:
    title = (lead.get("title") or "Untitled").strip()
    company = lead.get("company")
    source_url = lead.get("source_url")
    source_platform = lead.get("source_platform") or "unknown"
    location = lead.get("location")
    raw_text = lead.get("raw_text") or ""
    ats = extract_ats_identity(source_url)
    if ats and source_platform in {"unknown", "", "jobspy"}:
        source_platform = ats[0]
    meta = dict(lead.get("metadata") or {})
    if ats:
        meta["ats_platform"] = ats[0]
        meta["ats_org"] = ats[1]
        meta["ats_job_id"] = ats[2]
    work_model = (lead.get("work_model") or meta.get("work_model") or infer_work_model(location, raw_text))
    meta["work_model"] = work_model
    from app.modules.job_discovery.categories import classify_job

    source_cat = meta.get("category") if isinstance(meta.get("category"), str) else None
    classified = classify_job(title=title, raw_text=raw_text, source_category=source_cat)
    category = lead.get("category") or classified["category"]
    meta["categories"] = classified.get("categories") or []
    meta["category_label"] = classified.get("category_label")
    meta["category"] = category
    fp = listing_fingerprint(
        source_url=source_url,
        title=title,
        company=company,
        source_platform=source_platform,
    )
    return db.upsert_job_listing(
        fingerprint=fp,
        title=title,
        company=company,
        location=location,
        source_url=source_url,
        source_platform=source_platform,
        raw_text=raw_text,
        metadata=meta,
        status="active",
        work_model=str(work_model),
        category=str(category),
    )


async def ingest_queries(
    *,
    queries: list[str] | None = None,
    location: str | None = None,
    limit_per_query: int | None = None,
    sites: list[str] | None = None,
    hours_old: int | None = None,
    country_indeed: str = "USA",
    stale_after_hours: int | None = 24 * 21,
    quality_gate: bool | None = None,
) -> dict[str, Any]:
    """Write path: fan-out providers and upsert into job_listings."""
    q_list = [q.strip() for q in (queries or _default_queries()) if q.strip()]
    loc = location if location is not None else settings.JOB_INDEX_DEFAULT_LOCATION
    per = limit_per_query or settings.JOB_INDEX_INGEST_LIMIT
    freshness = hours_old if hours_old is not None else settings.JOB_INDEX_HOURS_OLD
    use_gate = settings.JOB_INDEX_QUALITY_GATE if quality_gate is None else quality_gate
    min_chars = int(settings.JOB_INDEX_MIN_JD_CHARS)
    created = 0
    updated = 0
    fetched = 0
    rejected = 0
    reject_reasons: dict[str, int] = {}
    errors: list[str] = []
    provider_rollup: dict[str, dict[str, Any]] = {}

    for query in q_list:
        per_query_stats: dict[str, Any] = {}
        try:
            leads = await discover_all(
                query=query,
                location=loc or None,
                limit=per,
                min_score=0.0,
                sites=sites,
                hours_old=freshness,
                country_indeed=country_indeed,
                user_id=None,
                skip_cache=True,
                provider_stats=per_query_stats,
            )
        except Exception as exc:  # noqa: BLE001 — ingest must continue
            errors.append(f"{query}: {exc}")
            continue
        fetched += len(leads)
        for name, info in per_query_stats.items():
            if name in {"unified_before_limit", "jobspy_sites", "cache_hit", "gather_error"}:
                continue
            if not isinstance(info, dict):
                continue
            bucket = provider_rollup.setdefault(name, {"count": 0, "errors": []})
            bucket["count"] += int(info.get("count") or 0)
            err = info.get("error")
            if err:
                bucket["errors"].append(f"{query}: {err}")

        if use_gate:
            leads, dropped = filter_quality_leads(leads, min_chars=min_chars)
            rejected += len(dropped)
            for d in dropped:
                reason = str(d.get("reject_reason") or "unknown")
                reject_reasons[reason] = reject_reasons.get(reason, 0) + 1

        for lead in leads:
            _, was_created = upsert_lead(lead)
            if was_created:
                created += 1
            else:
                updated += 1

    closed = 0
    if stale_after_hours and stale_after_hours > 0:
        closed = db.mark_stale_job_listings(max_age_hours=stale_after_hours)

    closed_seed = db.close_job_listings_by_platform("seed")
    closed_adzuna = db.close_job_listings_by_platform("adzuna")
    closed_thin = close_thin_active_listings(min_chars=min_chars) if use_gate else 0

    return {
        "queries": q_list,
        "location": loc or None,
        "hours_old": freshness,
        "limit_per_query": per,
        "quality_gate": use_gate,
        "fetched": fetched,
        "accepted": created + updated,
        "rejected_quality": rejected,
        "reject_reasons": reject_reasons,
        "created": created,
        "updated": updated,
        "closed_stale": closed,
        "closed_seed": closed_seed,
        "closed_adzuna": closed_adzuna,
        "closed_thin": closed_thin,
        "active_total": db.count_job_listings("active"),
        "provider_stats": provider_rollup,
        "errors": errors,
    }


def close_thin_active_listings(*, min_chars: int = 500) -> int:
    """Soft-close active rows that fail the real-JD quality gate (e.g. Adzuna teasers)."""
    closed = 0
    for row in db.search_job_listings(status="active", limit=500):
        verdict = assess_listing_quality(row, min_chars=min_chars)
        if verdict["ok"]:
            continue
        # Re-use platform closer one-by-one via status update
        with db.connect() as conn:
            conn.execute(
                "UPDATE job_listings SET status = 'closed', updated_at = ? WHERE id = ? AND status = 'active'",
                (db.utcnow(), row["id"]),
            )
        closed += 1
    return closed


def _default_queries() -> list[str]:
    from app.modules.job_discovery.categories import all_ingest_queries

    configured = [q.strip() for q in settings.JOB_INDEX_DEFAULT_QUERIES.split(",") if q.strip()]
    if not configured or (len(configured) == 1 and configured[0].lower() == "auto"):
        return all_ingest_queries()
    legacy = {"data analyst", "analytics", "business intelligence"}
    if set(q.lower() for q in configured) <= legacy:
        return all_ingest_queries()
    return configured


def search_index(
    *,
    query: str,
    location: str | None = None,
    limit: int = 10,
    min_score: float = 0.5,
    resume_text: str = "",
    max_age_hours: int | None = None,
    work_model: str | None = None,
    source_platform: str | None = None,
    category: str | None = None,
) -> list[dict[str, Any]]:
    """Read path: query local catalog, score, rank, hard-filter."""
    pool = db.search_job_listings(
        query=query if query and query.lower() not in {"*", "all"} else None,
        location=location,
        status="active",
        limit=max(limit * 5, 50),
        max_age_hours=max_age_hours if max_age_hours is not None else (hours_from_settings()),
        work_model=work_model,
        source_platform=source_platform,
        category=category,
    )
    if not pool and query.strip() and query.lower() not in {"*", "all"}:
        token = query.strip().split()[0]
        pool = db.search_job_listings(
            query=token,
            location=location,
            status="active",
            limit=max(limit * 5, 50),
            max_age_hours=max_age_hours if max_age_hours is not None else (hours_from_settings()),
            work_model=work_model,
            source_platform=source_platform,
            category=category,
        )

    from app.modules.job_discovery.scorer import score_job_detailed

    scored: list[dict[str, Any]] = []
    for item in pool:
        body = item.get("raw_text") or ""
        parsed_for_score = {
            "title": item.get("title") or "",
            "raw_text": body,
            "required_skills": [],
            "preferred_skills": [],
            "ats_keywords": [],
            "key_responsibilities": [],
        }
        detail = score_job_detailed(parsed_for_score, query, resume_text=resume_text)
        out = {
            "title": item.get("title"),
            "company": item.get("company"),
            "location": item.get("location"),
            "source_url": item.get("source_url"),
            "source_platform": item.get("source_platform") or "job_index",
            "raw_text": body,
            "metadata": {
                **(item.get("metadata") or {}),
                "listing_id": item.get("id"),
                "fingerprint": item.get("fingerprint"),
                "scraped_at": item.get("scraped_at"),
                "from_index": True,
                "score_breakdown": detail["score_breakdown"],
                "matched_skills": detail["matched_skills"],
                "missing_skills": detail["missing_skills"],
                "category": item.get("category"),
                "category_label": (item.get("metadata") or {}).get("category_label"),
            },
            "match_score": detail["match_score"],
            "scraped_at": item.get("scraped_at"),
            "category": item.get("category"),
        }
        scored.append(out)

    scored.sort(key=lambda x: x.get("match_score") or 0, reverse=True)
    threshold = min_score if min_score > 1 else min_score * 100
    high = [j for j in scored if (j.get("match_score") or 0) >= threshold]
    return (high or scored)[:limit]


def hours_from_settings() -> int | None:
    return 24 * 14


def _master_resume_plaintext() -> str:
    """Fallback profile text from the locked master DOCX when no uploaded resume exists."""
    try:
        from io import BytesIO

        from docx import Document

        from app.modules.resume_workspace.master_template import ensure_master_template_bytes

        data = ensure_master_template_bytes()
        if not data:
            return ""
        doc = Document(BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)
    except Exception:
        return ""


def resume_text_for_user(user_id: str | None) -> str:
    if user_id:
        latest = db.get_latest_resume(user_id)
        if latest:
            raw = latest.get("raw_text") or ""
            parsed = latest.get("parsed") or {}
            text = f"{raw} {' '.join(str(v) for v in parsed.values() if isinstance(v, str))}".strip()
            if text:
                return text
        template = db.get_active_template(user_id)
        if template and template.get("docx_bytes"):
            try:
                from io import BytesIO

                from docx import Document

                doc = Document(BytesIO(template["docx_bytes"]))
                text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip()).strip()
                if text:
                    return text
            except Exception:
                pass
    return _master_resume_plaintext()

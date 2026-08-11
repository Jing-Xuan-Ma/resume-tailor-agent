"""Module C: resume structure parsing — sections/entries/bullets, no hardcoded
counts, unmapped section titles surfaced instead of guessed."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from app.modules.resume_workspace.structure_parser import (
    CANONICAL_SECTIONS,
    parse_resume_structure,
)


def _section_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True


def _entry_heading(doc: Document, *lines: str) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _build_resume_a() -> bytes:
    """2 experience entries (2 + 3 bullets), 3 project entries (1 bullet each),
    2 education entries with no bullets at all."""
    doc = Document()

    _section_header(doc, "Professional Experience")
    _entry_heading(doc, "Software Engineer | Acme Corp", "Jun 2020 - Aug 2021")
    _bullet(doc, "Built scalable backend services in Python")
    _bullet(doc, "Reduced latency by optimizing database queries")

    _entry_heading(doc, "Data Analyst Intern | Beta Inc", "Jun 2019 - Aug 2019")
    _bullet(doc, "Analyzed customer churn data using SQL")
    _bullet(doc, "Built dashboards in Tableau for stakeholders")
    _bullet(doc, "Presented findings to leadership team")

    _section_header(doc, "Projects")
    for i in range(3):
        _entry_heading(doc, f"Side Project {i + 1} | Python, SQL")
        _bullet(doc, f"Implemented feature set number {i + 1}")

    _section_header(doc, "Education")
    _entry_heading(doc, "Johns Hopkins University")
    _entry_heading(doc, "Aug 2025 - Jun 2027")
    _entry_heading(doc, "University College Cork")
    _entry_heading(doc, "Sep 2021 - Jun 2025")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_resume_b() -> bytes:
    """3 experience entries (1 bullet each), 1 project entry (4 bullets) — a
    different shape from resume A, to prove nothing is hardcoded."""
    doc = Document()

    _section_header(doc, "Work Experience")
    for i in range(3):
        _entry_heading(doc, f"Role {i + 1} | Company {i + 1}", "2021 - 2022")
        _bullet(doc, f"Delivered outcome {i + 1} for the team")

    _section_header(doc, "Projects")
    _entry_heading(doc, "Flagship Project | Go, Kubernetes")
    for i in range(4):
        _bullet(doc, f"Shipped milestone {i + 1} on schedule")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_resume_a_structure_matches_actual_docx_shape() -> None:
    result = parse_resume_structure(_build_resume_a())
    sections = {s["type"]: s for s in result["sections"]}

    assert set(sections) == {"professional_experience", "projects", "education"}
    assert not result["unmapped_sections"]

    exp = sections["professional_experience"]
    assert len(exp["entries"]) == 2
    assert len(exp["entries"][0]["bullets"]) == 2
    assert len(exp["entries"][1]["bullets"]) == 3

    proj = sections["projects"]
    assert len(proj["entries"]) == 3
    assert all(len(e["bullets"]) == 1 for e in proj["entries"])

    edu = sections["education"]
    assert len(edu["entries"]) == 2
    assert edu["entries"][0]["title"] == "Johns Hopkins University"
    assert edu["entries"][0]["date_range"]
    assert edu["entries"][1]["title"] == "University College Cork"


def test_resume_b_has_a_completely_different_shape_and_still_parses() -> None:
    result = parse_resume_structure(_build_resume_b())
    sections = {s["type"]: s for s in result["sections"]}

    exp = sections["professional_experience"]
    assert len(exp["entries"]) == 3
    assert all(len(e["bullets"]) == 1 for e in exp["entries"])

    proj = sections["projects"]
    assert len(proj["entries"]) == 1
    assert len(proj["entries"][0]["bullets"]) == 4


def test_bullet_ids_and_verb_tense_are_populated() -> None:
    result = parse_resume_structure(_build_resume_a())
    exp = next(s for s in result["sections"] if s["type"] == "professional_experience")
    all_bullets = [b for e in exp["entries"] for b in e["bullets"]]
    ids = [b["id"] for b in all_bullets]
    assert len(ids) == len(set(ids)), "bullet ids must be unique"
    for bullet in all_bullets:
        assert bullet["text"], bullet
        assert bullet["verb_tense"] in {"past", "present", "present_participle"}
    # "Built ..." / "Reduced ..." both start with past-tense verbs.
    assert all_bullets[0]["verb_tense"] == "past"


def test_unrecognized_section_title_is_reported_not_guessed() -> None:
    doc = Document()
    _section_header(doc, "Volunteer Work")
    _entry_heading(doc, "Community Organizer | Local Nonprofit")
    _bullet(doc, "Coordinated weekend food drives")
    buf = BytesIO()
    doc.save(buf)

    result = parse_resume_structure(buf.getvalue())
    assert result["sections"] == []
    assert result["unmapped_sections"] == [{"raw_title": "VOLUNTEER WORK"}]


def test_custom_mapping_resolves_a_previously_unmapped_title() -> None:
    doc = Document()
    _section_header(doc, "Volunteer Work")
    _entry_heading(doc, "Community Organizer | Local Nonprofit")
    _bullet(doc, "Coordinated weekend food drives")
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    first = parse_resume_structure(docx_bytes)
    assert first["unmapped_sections"], "should be unmapped before a custom mapping exists"

    resolved = parse_resume_structure(
        docx_bytes, custom_mappings={"volunteer work": "competitions"}
    )
    assert resolved["unmapped_sections"] == []
    assert resolved["sections"][0]["type"] == "competitions"
    assert len(resolved["sections"][0]["entries"]) == 1


def test_canonical_section_registry_has_no_overlapping_phrases() -> None:
    seen: dict[str, str] = {}
    for section_type, phrases in CANONICAL_SECTIONS.items():
        for phrase in phrases:
            assert phrase not in seen, (
                f"'{phrase}' claimed by both {seen.get(phrase)} and {section_type}"
            )
            seen[phrase] = section_type

"""Module C: resume upload + version management API — .docx-only upload,
history retained across uploads (not overwritten), rollback to a previous
version, and the unmapped-section confirmation round trip."""

from __future__ import annotations

from io import BytesIO
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)


def _make_docx(section_title: str = "Professional Experience") -> bytes:
    doc = Document()
    header = doc.add_paragraph()
    run = header.add_run(section_title.upper())
    run.bold = True
    heading = doc.add_paragraph()
    heading_run = heading.add_run("Software Engineer | Acme Corp")
    heading_run.bold = True
    doc.add_paragraph("Jun 2020 - Aug 2021")
    doc.add_paragraph("Built scalable backend services", style="List Bullet")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_upload_rejects_non_docx_files() -> None:
    user = str(uuid4())
    resp = client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("resume.pdf", b"%PDF-1.4 not a real pdf", "application/pdf")},
    )
    assert resp.status_code == 400, resp.text


def test_upload_returns_parsed_resume_structure() -> None:
    user = str(uuid4())
    resp = client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("resume.docx", _make_docx(), "application/octet-stream")},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["is_active"] is True
    assert body["unmapped_sections"] == []
    sections = body["resume_structure"]["sections"]
    assert sections[0]["type"] == "professional_experience"
    entry = sections[0]["entries"][0]
    assert entry["title"] == "Software Engineer"
    assert entry["company"] == "Acme Corp"
    assert len(entry["bullets"]) == 1


def test_new_upload_keeps_history_instead_of_overwriting() -> None:
    user = str(uuid4())
    first = client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("v1.docx", _make_docx(), "application/octet-stream")},
    )
    second = client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("v2.docx", _make_docx(), "application/octet-stream")},
    )
    assert first.status_code == 200 and second.status_code == 200

    listed = client.get("/api/v1/resume-workspace/templates", params={"user_id": user})
    assert listed.status_code == 200, listed.text
    templates = listed.json()["templates"]
    assert len(templates) == 2, "both versions must remain queryable, not overwritten"

    ids = {t["id"] for t in templates}
    assert ids == {first.json()["template_id"], second.json()["template_id"]}

    active = [t for t in templates if t["is_active"]]
    assert len(active) == 1
    assert active[0]["id"] == second.json()["template_id"], "latest upload becomes current"


def test_listed_versions_carry_their_own_resume_structure() -> None:
    """Each version's card needs its own role/bullet counts, not the caller's
    globally-loaded profile inventory — so /templates must return the parsed
    structure per row, not just filename/is_active metadata."""
    user = str(uuid4())
    client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("v1.docx", _make_docx(), "application/octet-stream")},
    )
    listed = client.get("/api/v1/resume-workspace/templates", params={"user_id": user})
    templates = listed.json()["templates"]
    assert len(templates) == 1
    sections = templates[0]["resume_structure"]["sections"]
    assert sections[0]["type"] == "professional_experience"
    assert len(sections[0]["entries"]) == 1


def test_rollback_to_previous_version() -> None:
    user = str(uuid4())
    first = client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("v1.docx", _make_docx(), "application/octet-stream")},
    )
    client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("v2.docx", _make_docx(), "application/octet-stream")},
    )
    first_id = first.json()["template_id"]

    rollback = client.post(
        f"/api/v1/resume-workspace/template/{first_id}/activate",
        params={"user_id": user},
    )
    assert rollback.status_code == 200, rollback.text
    assert rollback.json()["ok"] is True

    active = client.get("/api/v1/resume-workspace/template/active", params={"user_id": user})
    assert active.status_code == 200
    assert active.json()["template_id"] == first_id, "rollback must switch the current version"

    listed = client.get("/api/v1/resume-workspace/templates", params={"user_id": user})
    templates = listed.json()["templates"]
    assert len(templates) == 2, "rollback must not delete the version rolled back from"


def test_activate_unknown_template_returns_404() -> None:
    user = str(uuid4())
    resp = client.post(
        "/api/v1/resume-workspace/template/nonexistent-id/activate",
        params={"user_id": user},
    )
    assert resp.status_code == 404


def test_unrecognized_section_surfaces_for_confirmation_not_silent_guess() -> None:
    user = str(uuid4())
    upload = client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("resume.docx", _make_docx("Volunteer Work"), "application/octet-stream")},
    )
    body = upload.json()
    assert body["resume_structure"].get("sections", []) == []
    assert body["unmapped_sections"] == [{"raw_title": "VOLUNTEER WORK"}]


def test_confirmed_section_mapping_is_saved_and_reused_on_reparse() -> None:
    user = str(uuid4())
    client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("resume.docx", _make_docx("Volunteer Work"), "application/octet-stream")},
    )

    mapped = client.post(
        "/api/v1/resume-workspace/template/section-mapping",
        json={"user_id": user, "raw_title": "VOLUNTEER WORK", "section_type": "competitions"},
    )
    assert mapped.status_code == 200, mapped.text
    body = mapped.json()
    assert body["unmapped_sections"] == []
    assert body["resume_structure"]["sections"][0]["type"] == "competitions"

    # A brand-new upload with the same non-standard heading should now resolve
    # automatically, without asking the user to confirm it again.
    second = client.post(
        "/api/v1/resume-workspace/template/upload",
        data={"user_id": user},
        files={"file": ("resume2.docx", _make_docx("Volunteer Work"), "application/octet-stream")},
    )
    assert second.json()["unmapped_sections"] == []
    assert second.json()["resume_structure"]["sections"][0]["type"] == "competitions"


def test_section_mapping_rejects_unknown_section_type() -> None:
    user = str(uuid4())
    resp = client.post(
        "/api/v1/resume-workspace/template/section-mapping",
        json={"user_id": user, "raw_title": "VOLUNTEER WORK", "section_type": "not_a_real_type"},
    )
    assert resp.status_code == 400
